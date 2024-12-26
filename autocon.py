import streamlit as st
import os
import time
import tempfile
from datetime import datetime
from typing import Tuple, Optional
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# PDF processing
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Document processing
from docx import Document

# OpenAI setup
from openai import OpenAI

# Initialize OpenAI client with API key from environment variable
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Template directory setup
TEMPLATE_DIR = "templates"
if not os.path.exists(TEMPLATE_DIR):
    os.makedirs(TEMPLATE_DIR)

# Global template cache
template_cache = {}

class ContractProcessor:
    @staticmethod
    @st.cache_resource
    def load_templates() -> dict:
        """Load all templates from the template directory."""
        try:
            if not os.path.exists(TEMPLATE_DIR):
                os.makedirs(TEMPLATE_DIR)
                default_template = ContractProcessor.generate_template("default")
                with open(os.path.join(TEMPLATE_DIR, "default.txt"), 'w') as file:
                    file.write(default_template)
                template_cache["default"] = default_template
            else:
                for filename in os.listdir(TEMPLATE_DIR):
                    if filename.endswith(".txt"):
                        template_name = filename[:-4]
                        with open(os.path.join(TEMPLATE_DIR, filename), 'r') as file:
                            template_cache[template_name] = file.read()
            return template_cache
        except Exception as e:
            st.error(f"Error loading templates: {str(e)}")
            return {}

    @staticmethod
    def generate_template(template_type: str) -> str:
        """Generate a new contract template using OpenAI."""
        try:
            prompt = f"Generate a {template_type} contract template. Include placeholders for specific details."
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal expert specializing in contract generation."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Error generating template: {str(e)}")
            return ""

    @staticmethod
    def extract_text_from_pdf(pdf_file) -> str:
        """Extract text from a PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
            return "\n".join(text)
        except Exception as e:
            st.error(f"Error extracting text from PDF: {str(e)}")
            return ""

    @staticmethod
    def read_file_content(uploaded_file) -> str:
        """Read content from various file types."""
        try:
            if uploaded_file.type == 'application/pdf':
                return ContractProcessor.extract_text_from_pdf(uploaded_file)
            else:
                return uploaded_file.getvalue().decode('utf-8')
        except Exception as e:
            st.error(f"Error reading file content: {str(e)}")
            return ""

    @staticmethod
    @st.cache_data
    def classify_template(text: str) -> str:
        """Classify the type of contract needed based on the proposal."""
        try:
            prompt = (
                f"Based on the following business proposal, what type of contract template "
                f"would be most appropriate? Please respond with only the contract type "
                f"name, nothing else.\n\n{text}\n\nContract type:"
            )
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal expert specializing in contract classification."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip().lower()
        except Exception as e:
            st.error(f"Error classifying template: {str(e)}")
            return "default"

    @staticmethod
    @st.cache_data
    def populate_template(template: str, proposal: str) -> str:
        """Populate the template with information from the proposal."""
        try:
            prompt = (
                f"Given the following contract template and business proposal, please fill "
                f"in the template with relevant information from the proposal:\n\n"
                f"Template:\n{template}\n\nProposal:\n{proposal}\n\nFilled template:"
            )
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal expert specializing in contract generation."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Error populating template: {str(e)}")
            return ""

    @staticmethod
    @st.cache_data
    def risk_compliance_check(contract: str) -> str:
        """Perform risk and compliance analysis on the contract."""
        try:
            prompt = (
                f"Analyze the following contract for potential risks and compliance "
                f"issues. Provide a detailed report:\n\n{contract}\n\n"
                f"Risk and Compliance Report:"
            )
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal expert specializing in risk assessment and compliance."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Error performing risk compliance check: {str(e)}")
            return ""

    @staticmethod
    @st.cache_data
    def generate_final_contract(contract: str) -> str:
        """Generate the final version of the contract."""
        try:
            prompt = (
                f"Review and finalize the following contract, ensuring consistency "
                f"and compliance:\n\n{contract}\n\nFinal Contract:"
            )
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal expert specializing in contract finalization."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Error generating final contract: {str(e)}")
            return ""

    @staticmethod
    def save_as_word(contract: str) -> str:
        """Save the contract as a Word document."""
        try:
            doc = Document()
            doc.add_heading('Legal Contract', 0)
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(contract)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc.save(tmp.name)
            return tmp.name
        except Exception as e:
            st.error(f"Error saving Word document: {str(e)}")
            return ""

    @staticmethod
    def save_as_pdf(contract: str) -> str:
        """Save the contract as a PDF document."""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                c = canvas.Canvas(tmp.name, pagesize=letter)
                width, height = letter
                y = height - 40
                
                # Add title and timestamp
                c.setFont("Helvetica-Bold", 16)
                c.drawString(40, y, "Legal Contract")
                y -= 30
                
                c.setFont("Helvetica", 10)
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                c.drawString(40, y, f"Generated on: {timestamp}")
                y -= 30
                
                # Add contract content
                c.setFont("Helvetica", 12)
                for line in contract.split('\n'):
                    if y < 40:
                        c.showPage()
                        y = height - 40
                    c.drawString(40, y, line)
                    y -= 15
                
                c.save()
            return tmp.name
        except Exception as e:
            st.error(f"Error saving PDF: {str(e)}")
            return ""

def process_proposal(proposal: str) -> Tuple[str, str, str, float]:
    """Process the business proposal and generate contract documents."""
    start_time = time.time()
    
    try:
        template_type = ContractProcessor.classify_template(proposal)
        template = ContractProcessor.generate_template(template_type)
        populated_contract = ContractProcessor.populate_template(template, proposal)
        risk_report = ContractProcessor.risk_compliance_check(populated_contract)
        
        processing_time = time.time() - start_time
        return template_type, populated_contract, risk_report, processing_time
    except Exception as e:
        st.error(f"Error processing proposal: {str(e)}")
        return "default", "", "", 0.0

def main():
    st.set_page_config(
        page_title="Legal Contract Automation System",
        page_icon="⚖️",
        layout="wide"
    )

    # Check if API key is configured
    if not os.getenv('OPENAI_API_KEY'):
        st.error("OpenAI API key not found. Please set OPENAI_API_KEY in your environment variables.")
        return

    st.title("⚖️ Legal Contract Automation System")
    
    # Initialize session state
    if 'populated_contract' not in st.session_state:
        st.session_state.populated_contract = None
    if 'initial_risk_report' not in st.session_state:
        st.session_state.initial_risk_report = None

    # Load templates
    ContractProcessor.load_templates()

    # Create main input area
    st.header("📝 Step 1: Input Business Proposal")
    
    input_type = st.radio(
        "Select input type:",
        ("Text", "File Upload"),
        horizontal=True
    )

    proposal = ""
    if input_type == "Text":
        proposal = st.text_area(
            "Enter your business proposal:",
            height=200,
            key="proposal_text"
        )
    else:
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'pdf'],
            key="file_uploader"
        )
        if uploaded_file is not None:
            with st.spinner("Reading file..."):
                proposal = ContractProcessor.read_file_content(uploaded_file)
                st.text_area("Extracted text:", proposal, height=200, key="extracted_text")

    process_button = st.button(
        "🚀 Process Proposal",
        key="process_button",
        use_container_width=True
    )

    if process_button and proposal:
        with st.spinner("Processing proposal..."):
            template_type, populated_contract, risk_report, processing_time = process_proposal(proposal)

        # Display results
        st.success(f"✓ Processed in {processing_time:.2f} seconds")
        st.info(f"📋 Suggested template type: {template_type}")

        st.session_state.populated_contract = populated_contract
        st.session_state.initial_risk_report = risk_report

    if st.session_state.populated_contract:
        # Display contract and risk report
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Generated Contract")
            reviewed_contract = st.text_area(
                "Review and edit the contract if necessary:",
                st.session_state.populated_contract,
                height=400,
                key="reviewed_contract"
            )

        with col2:
            st.subheader("⚠️ Risk and Compliance Report")
            st.text_area(
                "Risk Report",
                st.session_state.initial_risk_report,
                height=400,
                key="initial_risk_report",
                disabled=True
            )

        # Additional analysis
        if st.button("🔍 Run Additional Analysis", use_container_width=True):
            with st.spinner("Analyzing updated contract..."):
                additional_risk_report = ContractProcessor.risk_compliance_check(reviewed_contract)
                final_contract = ContractProcessor.generate_final_contract(reviewed_contract)

            st.subheader("📊 Additional Analysis Results")
            st.text_area(
                "Updated Risk Analysis",
                additional_risk_report,
                height=200
            )

            st.subheader("📋 Final Contract")
            st.text_area(
                "Final Version",
                final_contract,
                height=300,
                disabled=True
            )

            # Download options
            st.subheader("💾 Download Options")
            col3, col4 = st.columns(2)

            with col3:
                word_file = ContractProcessor.save_as_word(final_contract)
                if word_file:
                    with open(word_file, "rb") as file:
                        st.download_button(
                            label="📎 Download as Word",
                            data=file,
                            file_name=f"contract_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

            with col4:
                pdf_file = ContractProcessor.save_as_pdf(final_contract)
                if pdf_file:
                    with open(pdf_file, "rb") as file:
                        st.download_button(
                            label="📑 Download as PDF",
                            data=file,
                            file_name=f"contract_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )

            # Cleanup temporary files
            try:
                if word_file and os.path.exists(word_file):
                    os.unlink(word_file)
                if pdf_file and os.path.exists(pdf_file):
                    os.unlink(pdf_file)
            except Exception as e:
                st.warning(f"Warning: Could not clean up temporary files: {str(e)}")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        st.error("Please try refreshing the page or contact support if the issue persists.")